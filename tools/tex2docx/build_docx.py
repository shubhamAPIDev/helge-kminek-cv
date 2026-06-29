#!/usr/bin/env python3
"""Build an editable Word (.docx) version of the Helge Kminek CV.

Parses the arasgungore LaTeX source (main.tex header + sections/body.tex) and
renders it with python-docx, reproducing the PDF's layout: a centered header,
section titles underlined with a rule, and two-column entry rows with the date
right-aligned. The output stays plain editable Word text (borderless tables, no
text boxes), so Helge can edit it without touching LaTeX/Overleaf.

Self-contained: the only dependency is python-docx (see requirements.txt).

Usage:
    python3 tools/tex2docx/build_docx.py [output.docx]
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]


# --------------------------------------------------------------------------- #
# LaTeX text helpers
# --------------------------------------------------------------------------- #
def strip_comments(text: str) -> str:
    return re.sub(r"(?<!\\)%.*", "", text)


def extract_braced(text: str, pos: int) -> tuple[str, int]:
    """Given text[pos] == '{', return (inner, index after matching '}')."""
    depth, start = 0, pos
    while pos < len(text):
        ch = text[pos]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start + 1 : pos], pos + 1
        pos += 1
    return text[start + 1 :], len(text)


def get_args(text: str, pos: int, n: int) -> tuple[list[str], int]:
    """Read n consecutive {..} groups starting at/after pos."""
    args: list[str] = []
    while len(args) < n:
        while pos < len(text) and text[pos] in " \n\t":
            pos += 1
        if pos >= len(text) or text[pos] != "{":
            break
        inner, pos = extract_braced(text, pos)
        args.append(inner)
    return args, pos


def clean_text(text: str) -> str:
    """Strip LaTeX markup down to plain unicode text."""
    text = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = text.replace("\\%", "%").replace("\\&", "&").replace("\\$", "$")
    text = text.replace("\\#", "#").replace("\\_", "_")
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = text.replace(r"{\ss}", "ß").replace("\\ss", "ß")
    text = text.replace("---", "—").replace("--", "–")
    text = text.replace("\\,", " ").replace("~", " ")
    text = re.sub(r"\\[a-zA-Z]+\*?", "", text)  # drop remaining commands
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"[ \t]+", " ", text).strip()


# (text, bold, italic) run tuples
Run = tuple[str, bool, bool]


def parse_runs(text: str) -> list[Run]:
    """Split text into runs, honouring \\textbf{..} and \\textit{..}."""
    runs: list[Run] = []
    i = 0
    while i < len(text):
        m = re.compile(r"\\(textbf|textit|emph)\s*\{").search(text, i)
        if not m:
            tail = clean_text(text[i:])
            if tail:
                runs.append((tail, False, False))
            break
        if m.start() > i:
            head = clean_text(text[i : m.start()])
            if head:
                runs.append((head, False, False))
        inner, after = extract_braced(text, m.end() - 1)
        bold = m.group(1) == "textbf"
        italic = m.group(1) in ("textit", "emph")
        for t, b, it in parse_runs(inner):
            runs.append((t, b or bold, it or italic))
        i = after
    return [(t, b, it) for (t, b, it) in runs if t]


# --------------------------------------------------------------------------- #
# Parsing the CV into a structured model
# --------------------------------------------------------------------------- #
class Subheading:
    def __init__(self, org, date, role, loc):
        self.org, self.date, self.role, self.loc = org, date, role, loc
        self.bullets: list[list[Run]] = []


class SubHead:  # sub-subheading (e.g. semester label)
    def __init__(self, title):
        self.title = title


class Line:  # project heading / standalone bullet
    def __init__(self, runs, bullet=False):
        self.runs, self.bullet = runs, bullet


TOKEN = re.compile(
    r"\\(resumeSubheading|resumeSubSubheading|resumeProjectHeading|resumeItem|resumeSubItem)\b"
)


def parse_section(content: str) -> list:
    entries: list = []
    open_sub: Subheading | None = None
    pos = 0
    while True:
        m = TOKEN.search(content, pos)
        if not m:
            break
        kind = m.group(1)
        if kind == "resumeSubheading":
            args, pos = get_args(content, m.end(), 4)
            org, date, role, loc = (args + ["", "", "", ""])[:4]
            open_sub = Subheading(
                clean_text(org), clean_text(date), clean_text(role), clean_text(loc)
            )
            entries.append(open_sub)
        elif kind == "resumeSubSubheading":
            args, pos = get_args(content, m.end(), 2)
            entries.append(SubHead(clean_text(args[0]) if args else ""))
            open_sub = None
        elif kind == "resumeProjectHeading":
            args, pos = get_args(content, m.end(), 1)
            entries.append(Line(parse_runs(args[0] if args else "")))
            open_sub = None
        else:  # resumeItem / resumeSubItem
            args, pos = get_args(content, m.end(), 1)
            runs = parse_runs(args[0] if args else "")
            if open_sub is not None:
                open_sub.bullets.append(runs)
            else:
                entries.append(Line(runs, bullet=True))
    return entries


def parse_header(source: str) -> tuple[str, list[str]]:
    m = re.search(r"\\begin\{center\}(.*?)\\end\{center\}", source, re.DOTALL)
    if not m:
        return "", []
    hdr = m.group(1)
    hdr = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", hdr)
    hdr = re.sub(r"\\(faMapMarker|faAt)\b", "", hdr)
    hdr = re.sub(r"\\[hv]space\*?\{[^}]*\}", "", hdr)
    hdr = hdr.replace("$|$", "|").replace(r"{\ss}", "ß")
    lines: list[str] = []
    for part in re.split(r"\\\\", hdr):
        part = clean_text(part.replace("\\ ", " "))
        part = re.sub(r"\s*\|\s*", "   |   ", part).strip()
        if part:
            lines.append(part)
    return (lines[0], lines[1:]) if lines else ("", [])


def load_source() -> str:
    main = (REPO_ROOT / "main.tex").read_text(encoding="utf-8")
    body = (REPO_ROOT / "sections" / "body.tex").read_text(encoding="utf-8")
    return strip_comments(main.replace(r"\input{sections/body.tex}", body))


def parse_cv(source: str):
    name, contacts = parse_header(source)
    body_m = re.search(r"\\begin\{document\}(.*?)\\end\{document\}", source, re.DOTALL)
    body = body_m.group(1) if body_m else source
    sections = []
    heads = list(re.finditer(r"\\section\s*\{", body))
    for i, h in enumerate(heads):
        title, after = extract_braced(body, h.end() - 1)
        end = heads[i + 1].start() if i + 1 < len(heads) else len(body)
        entries = parse_section(body[after:end])
        if entries:
            sections.append((clean_text(title), entries))
    return name, contacts, sections


# --------------------------------------------------------------------------- #
# Rendering with python-docx
# --------------------------------------------------------------------------- #
def _no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _bottom_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)


def _strip_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _add_runs(paragraph, runs: list[Run], size=10):
    for text, bold, italic in runs:
        r = paragraph.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)


def _gap(doc, points: float):
    """Add a small empty paragraph of a given height (controls vertical spacing)."""
    p = doc.add_paragraph()
    _no_space(p)
    p.add_run("").font.size = Pt(points)


def _set_base_font(doc, name: str, size: float):
    """Set the default document font reliably (style font.name alone is flaky)."""
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def render(name, contacts, sections, output: Path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.85)
        sec.right_margin = Inches(0.85)
    # Serif body font (Times New Roman) to match the PDF's Computer Modern look.
    # Times is used because it is available on every Word install (Mac/Windows),
    # so the file renders identically for the recipient.
    _set_base_font(doc, "Times New Roman", 11)

    # Header: name (large small caps) + contact lines, centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name)
    r.font.size = Pt(22)
    r.font.small_caps = True
    for idx, line in enumerate(contacts):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(p)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.italic = idx == 0  # "Lebenslauf" subtitle

    for title, entries in sections:
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(12)
        head.paragraph_format.space_after = Pt(4)
        r = head.add_run(title)  # small caps render the section title like the PDF
        r.font.small_caps = True
        r.font.size = Pt(13)
        _bottom_rule(head)

        for e in entries:
            if isinstance(e, Subheading):
                _render_subheading(doc, e)
            elif isinstance(e, SubHead):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(e.title)
                r.bold = r.italic = True
                r.font.size = Pt(10)
            else:  # Line
                p = doc.add_paragraph(style="List Bullet" if e.bullet else None)
                if e.bullet:
                    p.paragraph_format.left_indent = Inches(0.25)
                _no_space(p)
                _add_runs(p, e.runs)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output}")


def _render_subheading(doc, e: Subheading):
    # One row, two cells. Left cell stacks org (bold) + role (italic); right cell
    # stacks date (bold) + location (italic), right-aligned. The date thus sits on
    # the same top line as the content, with no empty cells/rows to create gaps.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _strip_table_borders(table)
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.width = Inches(4.9)
    right.width = Inches(2.0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Left cell: org (bold), then role (italic) on its own line
    p = left.paragraphs[0]
    _no_space(p)
    first_used = False
    if e.org:
        r = p.add_run(e.org)
        r.bold = True
        r.font.size = Pt(11)
        first_used = True
    if e.role:
        rp = p if not first_used else left.add_paragraph()
        _no_space(rp)
        r = rp.add_run(e.role)
        r.italic = True
        r.font.size = Pt(10.5)

    # Right cell: date (regular, like the PDF), then location (italic), right-aligned
    p = right.paragraphs[0]
    _no_space(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if e.date:
        r = p.add_run(e.date)
        r.font.size = Pt(11)
    if e.loc:
        lp = right.add_paragraph()
        _no_space(lp)
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = lp.add_run(e.loc)
        r.italic = True
        r.font.size = Pt(10.5)

    for runs in e.bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.25)
        _no_space(bp)
        _add_runs(bp, runs)

    _gap(doc, 6)  # breathing room between entries


def main() -> int:
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else REPO_ROOT / "Helge_Kminek_Lebenslauf.docx"
    name, contacts, sections = parse_cv(load_source())
    render(name, contacts, sections, output)
    print(f"name={name!r}  contacts={len(contacts)}  sections={len(sections)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
